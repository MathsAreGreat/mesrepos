from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os


def create_school_report():
    # Create document
    doc = Document()

    # ========== Document Styles ==========
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ========== Title Page ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("\nRAPPORT D'ACTIVITÉ SCOLAIRE\n")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("École Secondaire Les Érables\n").bold = True

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}\n").italic = True

    doc.add_page_break()

    # ========== Activity Details ==========
    doc.add_heading("Atelier d'Écologie : Recyclage et Développement Durable", level=1)

    # 1. Objectives Section
    doc.add_heading("1. Objectifs de l'activité", level=2)
    objectives = [
        "Sensibiliser les élèves aux enjeux du recyclage",
        "Développer la créativité par des projets pratiques",
        "Promouvoir les comportements éco-responsables",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # 2. Participants Table - FIXED VERSION
    doc.add_heading("2. Participants", level=2)

    # Create table with proper dimensions
    participants_data = [
        ["Catégorie", "Détails"],
        ["Élèves impliqués", "Classes de 5ème A et B (30 élèves)"],
        ["Encadrants", "Mme Dupont (SVT), M. Martin (Arts Plastiques)"],
        ["Intervenants", "Association Éco-Action (2 animateurs)"],
        ["Durée", "2 heures - 14h00 à 16h00"],
    ]

    # Create table with correct number of rows
    participants = doc.add_table(rows=len(participants_data), cols=2)
    participants.alignment = WD_TABLE_ALIGNMENT.CENTER
    participants.style = "Light Shading Accent 1"

    # Populate table safely
    for i, row in enumerate(participants_data):
        for j, cell_value in enumerate(row):
            participants.cell(i, j).text = cell_value

    # 3. Activity Description
    doc.add_heading("3. Déroulement de l'activité", level=2)
    phases = [
        (
            "Introduction théorique (30 min)",
            "Présentation PowerPoint sur les enjeux environnementaux",
        ),
        ("Atelier pratique (1h)", "Création d'objets à partir de matériaux recyclés"),
        (
            "Restitution (30 min)",
            "Présentation des réalisations et évaluation collective",
        ),
    ]

    for phase, description in phases:
        p = doc.add_paragraph()
        p.add_run(phase + ": ").bold = True
        p.add_run(description)

    # 4. Results
    doc.add_heading("4. Résultats et Retours", level=2)
    doc.add_paragraph("Réalisations concrètes:", style="List Bullet")
    doc.add_paragraph("- 15 objets créés (stylos, porte-documents, décorations)")
    doc.add_paragraph("- 5 affiches éducatives réalisées")

    doc.add_paragraph("Retours des participants:", style="List Bullet")
    feedback = [
        ("Élèves", "90% ont trouvé l'activité 'intéressante' ou 'très intéressante'"),
        (
            "Enseignants",
            "Activité bien organisée avec un bon équilibre théorie/pratique",
        ),
        ("Intervenants", "Élèves engagés et créatifs"),
    ]

    for role, comment in feedback:
        p = doc.add_paragraph()
        p.add_run(f"{role}: ").bold = True
        p.add_run(comment)

    # 5. Conclusion
    doc.add_heading("5. Conclusion", level=2)
    conclusion = [
        "Activité globalement réussie avec une participation active des élèves",
        "Objectifs pédagogiques atteints",
        "À reconduire l'année prochaine avec éventuellement plus de temps pour la phase pratique",
    ]

    for point in conclusion:
        doc.add_paragraph(point, style="List Bullet")

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Rapport rédigé par:\n").italic = True
    p.add_run("Mme Dupont\n").bold = True
    p.add_run("Professeur de SVT").italic = True

    # Save document
    filename = f"Rapport_Activite_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(filename)
    print(f"Document généré avec succès: {os.path.abspath(filename)}")


if __name__ == "__main__":
    create_school_report()
