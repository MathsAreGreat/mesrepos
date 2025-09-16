from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Word document with improved formatting
doc = Document()

# Title (Centered, Bold, Large Font)
title = doc.add_paragraph()
title_run = title.add_run("Rapport d’Activité Scolaire")
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\n")  # Spacing

# General Information Section
doc.add_heading("Informations Générales", level=2)
info_fields = [
    "Titre de l’Activité", "Date", "Lieu",
    "Classe(s) concernée(s)", "Encadrants"
]
for field in info_fields:
    p = doc.add_paragraph()
    p.add_run(f"**{field} :** ").bold = True
    p.add_run("[...]")  # Placeholder for user input

doc.add_paragraph("\n")  # Spacing

# Introduction
doc.add_heading("1. Introduction", level=2)
intro_paragraph = doc.add_paragraph()
intro_text = ("Dans le cadre de [matière ou projet éducatif], une activité intitulée **[Nom de l’Activité]** "
              "a été organisée le **[date]** à **[lieu]**. Cette initiative visait à [objectif principal, "
              "ex. renforcer les connaissances des élèves sur un sujet spécifique, favoriser "
              "l’apprentissage par l’expérience, etc.].")
intro_paragraph.add_run(intro_text)

doc.add_paragraph("\n")  # Spacing

# Activity Details
doc.add_heading("2. Déroulement de l’Activité", level=2)
steps = [
    "**Présentation et préparation**\n- Explication des objectifs et des consignes aux élèves.\n"
    "- Répartition en groupes (si nécessaire).",
    "**Déroulement principal**\n- [Décrire les activités réalisées, ex. visite d’un site, expériences "
    "scientifiques, travaux de groupe, jeux éducatifs, etc.].\n- Interaction avec des intervenants (si applicable).",
    "**Clôture et retour d’expérience**\n- Débriefing avec les élèves pour partager leurs impressions.\n"
    "- Évaluation des acquis et discussion sur les apprentissages."
]
for step in steps:
    doc.add_paragraph(step, style="List Bullet")

doc.add_paragraph("\n")  # Spacing

# Results and Feedback
doc.add_heading("3. Résultats et Bilan", level=2)
feedback_fields = [
    "Points positifs", "Difficultés rencontrées", "Suggestions pour l’avenir"
]
for field in feedback_fields:
    p = doc.add_paragraph()
    p.add_run(f"**{field} :** ").bold = True
    p.add_run("[...]")

doc.add_paragraph("\n")  # Spacing

# Conclusion
doc.add_heading("4. Conclusion", level=2)
conclusion_paragraph = doc.add_paragraph()
conclusion_text = ("L’activité **[Nom de l’Activité]** a été une expérience enrichissante pour les élèves et leur "
                   "a permis de [résumé des bénéfices]. Grâce à cette initiative, ils ont pu **[mentionner les "
                   "compétences ou savoirs acquis]**. Il serait pertinent de renouveler ce type d’événement pour "
                   "**[suggestions pour l’avenir]**.")
conclusion_paragraph.add_run(conclusion_text)

doc.add_paragraph("\n")  # Spacing

# Signature Section (Centered)
signature = doc.add_paragraph()
signature_run = signature.add_run(
    "\nFait à [Lieu], le [Date]\n[Nom et Signature de l’enseignant(e) responsable]")
signature_run.font.size = Pt(12)
signature.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Save the improved document
better_file_path = "rapport_activite_scolaire_better.docx"
doc.save(better_file_path)
